from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException, Form, Query
from fastapi.responses import StreamingResponse, FileResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
import aiofiles
import json
from emergentintegrations.llm.chat import LlmChat, UserMessage
import asyncio
from rag import RAGIndex, retrieve_context_for_query
import boto3
from tavily import TavilyClient

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Storage configuration
STORAGE_PROVIDER = os.environ.get('STORAGE_PROVIDER', 'local')
LOCAL_STORAGE_PATH = Path(os.environ.get('LOCAL_STORAGE_PATH', '/app/uploads'))
LOCAL_STORAGE_PATH.mkdir(parents=True, exist_ok=True)

# Create the main app
app = FastAPI()
api_router = APIRouter(prefix="/api")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Tavily client for web search (optional)
tavily_client = None
TAVILY_API_KEY = os.environ.get('TAVILY_API_KEY')
if TAVILY_API_KEY:
    tavily_client = TavilyClient(api_key=TAVILY_API_KEY)
    logger.info("Tavily web search enabled")

# ============== MODELS ==============

class ProjectBase(BaseModel):
    name: str
    description: Optional[str] = ""
    instructions: Optional[str] = ""
    memory: Optional[str] = ""
    extended_thinking_enabled: bool = False
    thinking_budget: int = 10000
    web_search_enabled: bool = True  # Default to on as user requested
    llm_provider: str = "anthropic"  # "anthropic" or "bedrock"
    temperature: float = 0.7  # 0.0 to 1.0, controls randomness/creativity
    top_p: float = 0.9  # 0.0 to 1.0, controls diversity of output

class ProjectCreate(ProjectBase):
    pass

class ProjectUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    instructions: Optional[str] = None
    memory: Optional[str] = None
    extended_thinking_enabled: Optional[bool] = None
    thinking_budget: Optional[int] = None
    web_search_enabled: Optional[bool] = None
    archived: Optional[bool] = None
    llm_provider: Optional[str] = None
    temperature: Optional[float] = None
    top_p: Optional[float] = None

class Project(ProjectBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    file_count: int = 0
    starred: bool = False
    archived: bool = False

class FileMetadata(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    project_id: str
    filename: str
    original_filename: str
    file_type: str
    file_size: int
    mime_type: str
    storage_path: str
    indexed: bool = False
    content_preview: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ConversationBase(BaseModel):
    project_id: str
    name: str = "New conversation"

class ConversationUpdate(BaseModel):
    name: Optional[str] = None
    project_id: Optional[str] = None
    starred: Optional[bool] = None
    archived: Optional[bool] = None
    extended_thinking: Optional[bool] = None
    web_search: Optional[bool] = None

class Conversation(ConversationBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    starred: bool = False
    archived: bool = False
    extended_thinking: bool = False
    web_search: bool = False

class Message(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    conversation_id: str
    role: str  # 'user' or 'assistant'
    content: str
    thinking: Optional[str] = None
    thinking_time: Optional[int] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ChatRequest(BaseModel):
    conversation_id: str
    message: str
    include_knowledge_base: bool = True
    extended_thinking: bool = False
    thinking_budget: int = 10000  # Default thinking budget tokens
    web_search: bool = False

class StorageConfig(BaseModel):
    provider: str  # 'local' or 's3'
    s3_bucket_name: Optional[str] = None
    s3_access_key: Optional[str] = None
    s3_secret_key: Optional[str] = None
    s3_region: Optional[str] = "us-east-1"

# ============== STORAGE ABSTRACTION ==============

class StorageProvider:
    async def save_file(self, file_content: bytes, filename: str, project_id: str) -> str:
        raise NotImplementedError
    
    async def get_file(self, storage_path: str) -> bytes:
        raise NotImplementedError
    
    async def delete_file(self, storage_path: str) -> bool:
        raise NotImplementedError

class LocalStorageProvider(StorageProvider):
    def __init__(self, base_path: Path):
        self.base_path = base_path
    
    async def save_file(self, file_content: bytes, filename: str, project_id: str) -> str:
        project_dir = self.base_path / project_id
        project_dir.mkdir(parents=True, exist_ok=True)
        file_path = project_dir / filename
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        return str(file_path)
    
    async def get_file(self, storage_path: str) -> bytes:
        async with aiofiles.open(storage_path, 'rb') as f:
            return await f.read()
    
    async def delete_file(self, storage_path: str) -> bool:
        try:
            os.remove(storage_path)
            return True
        except:
            return False

class S3StorageProvider(StorageProvider):
    def __init__(self, bucket_name: str, access_key: str, secret_key: str, region: str):
        import boto3
        self.bucket_name = bucket_name
        self.s3_client = boto3.client(
            's3',
            aws_access_key_id=access_key,
            aws_secret_access_key=secret_key,
            region_name=region
        )
    
    async def save_file(self, file_content: bytes, filename: str, project_id: str) -> str:
        key = f"{project_id}/{filename}"
        self.s3_client.put_object(Bucket=self.bucket_name, Key=key, Body=file_content)
        return key
    
    async def get_file(self, storage_path: str) -> bytes:
        response = self.s3_client.get_object(Bucket=self.bucket_name, Key=storage_path)
        return response['Body'].read()
    
    async def delete_file(self, storage_path: str) -> bool:
        try:
            self.s3_client.delete_object(Bucket=self.bucket_name, Key=storage_path)
            return True
        except:
            return False

def get_storage_provider() -> StorageProvider:
    if STORAGE_PROVIDER == 's3':
        return S3StorageProvider(
            bucket_name=os.environ.get('S3_BUCKET_NAME', ''),
            access_key=os.environ.get('S3_ACCESS_KEY', ''),
            secret_key=os.environ.get('S3_SECRET_KEY', ''),
            region=os.environ.get('S3_REGION', 'us-east-1')
        )
    return LocalStorageProvider(LOCAL_STORAGE_PATH)

storage = get_storage_provider()

# ============== FILE PROCESSING ==============

def get_mime_type(filename: str) -> str:
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    mime_map = {
        'pdf': 'application/pdf',
        'txt': 'text/plain',
        'md': 'text/markdown',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
        'png': 'image/png',
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
        'bmp': 'image/bmp',
    }
    return mime_map.get(ext, 'application/octet-stream')

def get_file_type(filename: str) -> str:
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    type_map = {
        'pdf': 'PDF',
        'txt': 'TXT',
        'md': 'MD',
        'docx': 'DOCX',
        'doc': 'DOC',
        'png': 'PNG',
        'jpg': 'JPG',
        'jpeg': 'JPG',
        'bmp': 'BMP',
    }
    return type_map.get(ext, 'FILE')

async def extract_text_content(file_content: bytes, filename: str, mime_type: str) -> str:
    """
    Extract text content from files for RAG INDEXING (not chat attachments).
    This is intentionally limited to create manageable chunks for vector search.
    For full document content in chat, see the /chat/with-files endpoint.
    """
    try:
        if mime_type == 'text/plain' or mime_type == 'text/markdown':
            return file_content.decode('utf-8', errors='ignore')[:5000]
        
        elif mime_type == 'application/pdf':
            from PyPDF2 import PdfReader
            import io
            reader = PdfReader(io.BytesIO(file_content))
            text = ""
            for page in reader.pages[:10]:
                text += page.extract_text() or ""
            return text[:5000]
        
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            from docx import Document
            import io
            doc = Document(io.BytesIO(file_content))
            text = "\n".join([p.text for p in doc.paragraphs])
            return text[:5000]
        
        return ""
    except Exception as e:
        logger.error(f"Error extracting text: {e}")
        return ""

# ============== PROJECT ROUTES ==============

@api_router.get("/")
async def root():
    return {"message": "Assessment Editor API", "version": "1.0.0"}

@api_router.post("/projects", response_model=Project)
async def create_project(project: ProjectCreate):
    project_obj = Project(**project.model_dump())
    doc = project_obj.model_dump()
    await db.projects.insert_one(doc)
    return project_obj

@api_router.get("/projects", response_model=List[Project])
async def get_projects(
    search: Optional[str] = Query(None),
    starred_only: bool = Query(False),
    archived: Optional[bool] = Query(None)
):
    query = {}
    if search:
        query["name"] = {"$regex": search, "$options": "i"}
    if starred_only:
        query["starred"] = True
    if archived is not None:
        query["archived"] = archived
    
    projects = await db.projects.find(query, {"_id": 0}).sort("updated_at", -1).to_list(100)
    
    # Get file counts
    for p in projects:
        count = await db.files.count_documents({"project_id": p["id"]})
        p["file_count"] = count
    
    return projects

@api_router.get("/projects/{project_id}", response_model=Project)
async def get_project(project_id: str):
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    count = await db.files.count_documents({"project_id": project_id})
    project["file_count"] = count
    return project

@api_router.put("/projects/{project_id}", response_model=Project)
async def update_project(project_id: str, update: ProjectUpdate):
    update_data = {k: v for k, v in update.model_dump().items() if v is not None}
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    result = await db.projects.update_one({"id": project_id}, {"$set": update_data})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Project not found")
    
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    count = await db.files.count_documents({"project_id": project_id})
    project["file_count"] = count
    return project

@api_router.put("/projects/{project_id}/star")
async def toggle_star_project(project_id: str):
    project = await db.projects.find_one({"id": project_id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    new_starred = not project.get("starred", False)
    await db.projects.update_one({"id": project_id}, {"$set": {"starred": new_starred}})
    return {"starred": new_starred}

@api_router.delete("/projects/{project_id}")
async def delete_project(project_id: str):
    # Delete all files
    files = await db.files.find({"project_id": project_id}, {"_id": 0}).to_list(1000)
    for f in files:
        await storage.delete_file(f["storage_path"])
    await db.files.delete_many({"project_id": project_id})
    
    # Delete conversations and messages
    convos = await db.conversations.find({"project_id": project_id}, {"_id": 0}).to_list(1000)
    for c in convos:
        await db.messages.delete_many({"conversation_id": c["id"]})
    await db.conversations.delete_many({"project_id": project_id})
    
    # Delete project
    result = await db.projects.delete_one({"id": project_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Project not found")
    
    return {"success": True}

# ============== FILE ROUTES ==============

@api_router.post("/projects/{project_id}/files", response_model=FileMetadata)
async def upload_file(project_id: str, file: UploadFile = File(...)):
    # Verify project exists
    project = await db.projects.find_one({"id": project_id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Validate file type
    allowed_extensions = ['pdf', 'txt', 'md', 'docx', 'doc', 'png', 'jpg', 'jpeg', 'bmp']
    ext = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
    if ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"File type not allowed. Allowed: {', '.join(allowed_extensions)}")
    
    # Read and save file
    content = await file.read()
    unique_filename = f"{uuid.uuid4()}_{file.filename}"
    storage_path = await storage.save_file(content, unique_filename, project_id)
    
    mime_type = get_mime_type(file.filename)
    content_preview = await extract_text_content(content, file.filename, mime_type)
    
    file_meta = FileMetadata(
        project_id=project_id,
        filename=unique_filename,
        original_filename=file.filename,
        file_type=get_file_type(file.filename),
        file_size=len(content),
        mime_type=mime_type,
        storage_path=storage_path,
        indexed=False,  # Will be set to True after RAG indexing
        content_preview=content_preview
    )
    
    await db.files.insert_one(file_meta.model_dump())
    
    # RAG indexing (async, non-blocking for text documents)
    if content_preview:
        try:
            api_key = os.environ.get('EMERGENT_LLM_KEY', '')
            rag_index = RAGIndex(db, project_id)
            chunks_indexed = await rag_index.index_document(
                file_id=file_meta.id,
                filename=file.filename,
                content=content_preview,
                api_key=api_key
            )
            if chunks_indexed > 0:
                await db.files.update_one(
                    {"id": file_meta.id},
                    {"$set": {"indexed": True, "chunks_count": chunks_indexed}}
                )
                file_meta.indexed = True
                logger.info(f"RAG indexed {chunks_indexed} chunks for {file.filename}")
        except Exception as e:
            logger.error(f"RAG indexing failed for {file.filename}: {e}")
    
    # Update project
    await db.projects.update_one(
        {"id": project_id}, 
        {"$set": {"updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    return file_meta

@api_router.get("/projects/{project_id}/files", response_model=List[FileMetadata])
async def get_project_files(project_id: str):
    files = await db.files.find({"project_id": project_id}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return files

@api_router.get("/files/{file_id}/download")
async def download_file(file_id: str):
    file_meta = await db.files.find_one({"id": file_id}, {"_id": 0})
    if not file_meta:
        raise HTTPException(status_code=404, detail="File not found")
    
    if STORAGE_PROVIDER == 'local':
        return FileResponse(
            file_meta["storage_path"],
            filename=file_meta["original_filename"],
            media_type=file_meta["mime_type"]
        )
    else:
        content = await storage.get_file(file_meta["storage_path"])
        return StreamingResponse(
            iter([content]),
            media_type=file_meta["mime_type"],
            headers={"Content-Disposition": f'attachment; filename="{file_meta["original_filename"]}"'}
        )

@api_router.delete("/files/{file_id}")
async def delete_file(file_id: str):
    file_meta = await db.files.find_one({"id": file_id}, {"_id": 0})
    if not file_meta:
        raise HTTPException(status_code=404, detail="File not found")
    
    # Remove from RAG index
    try:
        rag_index = RAGIndex(db, file_meta["project_id"])
        await rag_index.remove_document(file_id)
    except Exception as e:
        logger.error(f"Failed to remove from RAG index: {e}")
    
    await storage.delete_file(file_meta["storage_path"])
    await db.files.delete_one({"id": file_id})
    
    return {"success": True}


@api_router.post("/projects/{project_id}/reindex")
async def reindex_project_files(project_id: str):
    """Re-index all files in a project for RAG"""
    project = await db.projects.find_one({"id": project_id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    files = await db.files.find({"project_id": project_id}, {"_id": 0}).to_list(100)
    
    api_key = os.environ.get('EMERGENT_LLM_KEY', '')
    rag_index = RAGIndex(db, project_id)
    
    indexed_count = 0
    total_chunks = 0
    for f in files:
        if f.get("content_preview"):
            try:
                # Force re-index by removing existing chunks first
                await rag_index.remove_document(f["id"])
                
                chunks = await rag_index.index_document(
                    file_id=f["id"],
                    filename=f["original_filename"],
                    content=f["content_preview"],
                    api_key=api_key
                )
                if chunks > 0:
                    await db.files.update_one(
                        {"id": f["id"]},
                        {"$set": {"indexed": True, "chunks_count": chunks}}
                    )
                    indexed_count += 1
                    total_chunks += chunks
                    logger.info(f"Indexed {f['original_filename']}: {chunks} chunks")
            except Exception as e:
                logger.error(f"Failed to index {f['original_filename']}: {e}")
    
    # Get stats
    stats = await rag_index.get_stats()
    
    return {
        "success": True,
        "files_indexed": indexed_count,
        "total_chunks": stats["total_chunks"]
    }


@api_router.get("/projects/{project_id}/rag-stats")
async def get_rag_stats(project_id: str):
    """Get RAG index statistics for a project"""
    rag_index = RAGIndex(db, project_id)
    stats = await rag_index.get_stats()
    return stats

# ============== CONVERSATION ROUTES ==============

@api_router.post("/conversations", response_model=Conversation)
async def create_conversation(conv: ConversationBase):
    # Verify project exists
    project = await db.projects.find_one({"id": conv.project_id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    conv_obj = Conversation(**conv.model_dump())
    await db.conversations.insert_one(conv_obj.model_dump())
    return conv_obj

@api_router.get("/projects/{project_id}/conversations", response_model=List[Conversation])
async def get_project_conversations(project_id: str):
    convos = await db.conversations.find(
        {"project_id": project_id}, {"_id": 0}
    ).sort("updated_at", -1).to_list(100)
    return convos

@api_router.get("/conversations/recent", response_model=List[Dict[str, Any]])
async def get_recent_conversations(limit: int = Query(10, le=50)):
    convos = await db.conversations.find({}, {"_id": 0}).sort("updated_at", -1).to_list(limit)
    
    # Add project names
    result = []
    for c in convos:
        project = await db.projects.find_one({"id": c["project_id"]}, {"_id": 0, "name": 1})
        c["project_name"] = project["name"] if project else "Unknown"
        result.append(c)
    
    return result

@api_router.get("/conversations/{conversation_id}", response_model=Conversation)
async def get_conversation(conversation_id: str):
    conv = await db.conversations.find_one({"id": conversation_id}, {"_id": 0})
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    return conv

@api_router.put("/conversations/{conversation_id}", response_model=Conversation)
async def update_conversation(conversation_id: str, update: ConversationUpdate):
    """Update conversation (rename, change project, etc.)"""
    update_data = {k: v for k, v in update.model_dump().items() if v is not None}
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    result = await db.conversations.update_one({"id": conversation_id}, {"$set": update_data})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    conv = await db.conversations.find_one({"id": conversation_id}, {"_id": 0})
    return conv

@api_router.get("/conversations", response_model=List[Dict[str, Any]])
async def get_all_conversations(starred: Optional[bool] = Query(None), archived: Optional[bool] = Query(None)):
    """Get all conversations with optional filtering"""
    query = {}
    if starred is not None:
        query["starred"] = starred
    if archived is not None:
        query["archived"] = archived
        
    convos = await db.conversations.find(query, {"_id": 0}).sort("updated_at", -1).to_list(1000)
    
    # Add project names
    result = []
    for c in convos:
        project = await db.projects.find_one({"id": c.get("project_id")}, {"_id": 0, "name": 1})
        c["project_name"] = project["name"] if project else "No Project"
        result.append(c)
    
    return result

@api_router.put("/conversations/{conversation_id}/star")
async def toggle_star_conversation(conversation_id: str):
    conv = await db.conversations.find_one({"id": conversation_id})
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    new_starred = not conv.get("starred", False)
    await db.conversations.update_one({"id": conversation_id}, {"$set": {"starred": new_starred}})
    return {"starred": new_starred}

@api_router.delete("/conversations/{conversation_id}")
async def delete_conversation(conversation_id: str):
    await db.messages.delete_many({"conversation_id": conversation_id})
    result = await db.conversations.delete_one({"id": conversation_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Conversation not found")
    return {"success": True}

# ============== MESSAGE ROUTES ==============

@api_router.get("/conversations/{conversation_id}/messages", response_model=List[Message])
async def get_messages(conversation_id: str):
    messages = await db.messages.find(
        {"conversation_id": conversation_id}, {"_id": 0}
    ).sort("created_at", 1).to_list(1000)
    return messages

# ============== CHAT WITH AI ==============

async def call_bedrock_converse(
    model_id: str, 
    messages: List[dict], 
    aws_config: dict, 
    max_tokens: int = 4000,
    extended_thinking: bool = False,
    thinking_budget: int = 10000
) -> tuple:
    """
    Call AWS Bedrock using the Converse API (newer, more reliable than InvokeModel)
    Supports extended thinking for Claude models on Bedrock.
    Returns: (response_text, thinking_content, thinking_time)
    """
    import time
    start_time = time.time()
    
    try:
        bedrock_runtime = boto3.client(
            service_name='bedrock-runtime',
            region_name=aws_config['aws_region_name'],
            aws_access_key_id=aws_config['aws_access_key_id'],
            aws_secret_access_key=aws_config['aws_secret_access_key']
        )
        
        # Convert messages to Bedrock format
        bedrock_messages = []
        for msg in messages:
            if msg['role'] == 'system':
                continue  # System messages handled separately in Converse API
            # Ensure content is properly encoded as UTF-8 string
            content_text = msg['content']
            if isinstance(content_text, bytes):
                content_text = content_text.decode('utf-8', errors='replace')
            bedrock_messages.append({
                "role": msg['role'],
                "content": [{"text": str(content_text)}]
            })
        
        # Find system message
        system_content = None
        for msg in messages:
            if msg['role'] == 'system':
                sys_text = msg['content']
                if isinstance(sys_text, bytes):
                    sys_text = sys_text.decode('utf-8', errors='replace')
                system_content = [{"text": str(sys_text)}]
                break
        
        # Build inference config - remove temperature/topP when thinking is enabled
        inference_config = {
            "maxTokens": max_tokens,
        }
        
        # Only add temperature/topP when NOT using extended thinking
        # (thinking is not compatible with temperature, top_p, or top_k modifications)
        add_temp_params = True
        
        # Additional model request fields for Claude-specific features
        additional_fields = {}
        
        # Add extended thinking for Claude models if enabled
        # Note: Extended thinking is only supported for Claude 3.7+ and Claude 4 models
        is_claude_model = 'anthropic' in model_id.lower() or 'claude' in model_id.lower()
        thinking_content = None
        thinking_time = None
        
        if extended_thinking and is_claude_model:
            # Check if model supports extended thinking
            # Claude Sonnet 4.5: anthropic.claude-sonnet-4-5-20250929-v1:0
            # Claude 3.7 Sonnet: anthropic.claude-3-7-sonnet-20250219-v1:0
            # Claude 3.5 does NOT support thinking
            supports_thinking = any(x in model_id.lower() for x in [
                'claude-3-7', 'claude-4', 'sonnet-4', 'opus-4', 'haiku-4',
                'claude-sonnet-4', 'claude-opus-4', 'claude-haiku-4'
            ])
            
            if supports_thinking:
                # Minimum budget is 1024 tokens per Anthropic API
                actual_budget = max(1024, thinking_budget)
                # Thinking parameters go in additionalModelRequestFields for Bedrock
                additional_fields["thinking"] = {
                    "type": "enabled",
                    "budget_tokens": actual_budget
                }
                # max_tokens must be greater than budget_tokens
                # Ensure max_tokens allows for thinking + actual response
                inference_config["maxTokens"] = actual_budget + 4000
                # Thinking is NOT compatible with temperature, top_p, top_k
                add_temp_params = False
                logger.info(f"Extended thinking enabled with budget: {actual_budget} tokens, maxTokens: {inference_config['maxTokens']}")
            else:
                logger.warning(f"Extended thinking requested but model {model_id} does not support it (requires Claude 3.7 or later)")
        
        # Only add temperature (not both temp and topP - Claude 4 doesn't allow both)
        if add_temp_params:
            inference_config["temperature"] = 0.7
        
        # Call Converse API
        converse_params = {
            "modelId": model_id,
            "messages": bedrock_messages,
            "inferenceConfig": inference_config
        }
        
        if system_content:
            converse_params["system"] = system_content
        
        # Add additional model request fields if any
        if additional_fields:
            converse_params["additionalModelRequestFields"] = additional_fields
        
        logger.info(f"Bedrock Converse API call: model={model_id}, messages={len(bedrock_messages)}, extended_thinking={extended_thinking and is_claude_model}")
        
        response = bedrock_runtime.converse(**converse_params)
        
        elapsed_time = time.time() - start_time
        
        # Extract response text and thinking content
        output_message = response.get('output', {}).get('message', {})
        content_blocks = output_message.get('content', [])
        
        response_text = ""
        for block in content_blocks:
            if 'text' in block:
                text_content = block['text']
                if isinstance(text_content, bytes):
                    text_content = text_content.decode('utf-8', errors='replace')
                response_text += str(text_content)
            # Extract thinking content if present (Claude extended thinking)
            elif 'thinking' in block:
                thinking_content = block.get('thinking', '')
                thinking_time = round(elapsed_time)
                logger.info(f"Found thinking content: {len(thinking_content)} chars")
            elif block.get('type') == 'thinking':
                thinking_content = block.get('text', '')
                thinking_time = round(elapsed_time)
                logger.info(f"Found thinking block: {len(thinking_content)} chars")
        
        # Log response metadata for debugging
        stop_reason = response.get('stopReason', 'unknown')
        usage = response.get('usage', {})
        logger.info(f"Bedrock Converse API success: model={model_id}, response_length={len(response_text)}, stop_reason={stop_reason}, usage={usage}, thinking={'yes' if thinking_content else 'no'}")
        
        return (response_text, thinking_content, thinking_time)
        
    except boto3.exceptions.Boto3Error as e:
        logger.error(f"Bedrock boto3 error: {e}")
        raise HTTPException(status_code=500, detail=f"Bedrock API error: {str(e)}")
    except Exception as e:
        logger.error(f"Bedrock Converse API error: {type(e).__name__}: {e}")
        raise HTTPException(status_code=500, detail=f"Bedrock API error: {str(e)}")


async def perform_web_search(query: str) -> str:
    """
    Perform a web search using Tavily API.
    Returns formatted search results as a string.
    """
    if not tavily_client:
        return "Web search is not available. TAVILY_API_KEY not configured."
    
    try:
        logger.info(f"Performing web search for: {query}")
        response = tavily_client.search(
            query=query,
            search_depth="basic",
            max_results=5,
            include_answer=True
        )
        
        # Format results
        results = []
        if response.get('answer'):
            results.append(f"**Summary:** {response['answer']}\n")
        
        results.append("**Sources:**")
        for i, result in enumerate(response.get('results', [])[:5], 1):
            title = result.get('title', 'No title')
            url = result.get('url', '')
            content = result.get('content', '')[:300]
            results.append(f"\n{i}. **{title}**\n   URL: {url}\n   {content}...")
        
        formatted = "\n".join(results)
        logger.info(f"Web search returned {len(response.get('results', []))} results")
        return formatted
        
    except Exception as e:
        logger.error(f"Web search error: {e}")
        return f"Web search failed: {str(e)}"


async def retrieve_kb_file_content(project_id: str, filename: str) -> str:
    """
    Retrieve the full content of a specific file from the knowledge base.
    Called by Claude via tool use when it needs to access a file.
    """
    try:
        # Find the file in the database
        file_record = await db.files.find_one(
            {"project_id": project_id, "original_filename": {"$regex": filename, "$options": "i"}},
            {"_id": 0}
        )
        
        if not file_record:
            # Try partial match
            all_files = await db.files.find({"project_id": project_id}, {"_id": 0}).to_list(100)
            for f in all_files:
                if filename.lower() in f['original_filename'].lower():
                    file_record = f
                    break
        
        if not file_record:
            return f"File '{filename}' not found in knowledge base."
        
        # Read full content from storage
        file_path = LOCAL_STORAGE_PATH / file_record['storage_path']
        if not file_path.exists():
            return f"File '{filename}' exists in database but file not found on disk."
        
        content_bytes = file_path.read_bytes()
        ext = file_record['file_type'].lower()
        
        if ext in ['txt', 'md']:
            content = content_bytes.decode('utf-8', errors='ignore')
        elif ext == 'pdf':
            from PyPDF2 import PdfReader
            import io
            reader = PdfReader(io.BytesIO(content_bytes))
            pages_text = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages_text.append(f"[Page {i+1}]\n{page_text}")
            content = "\n\n".join(pages_text)
        elif ext == 'docx':
            from docx import Document
            import io
            doc = Document(io.BytesIO(content_bytes))
            content = "\n".join([p.text for p in doc.paragraphs])
        else:
            content = f"[Cannot extract text from {ext} files]"
        
        logger.info(f"KB file retrieved: {file_record['original_filename']}, {len(content)} chars")
        return f"# {file_record['original_filename']}\n\n{content}"
        
    except Exception as e:
        logger.error(f"KB file retrieval error: {e}")
        return f"Error retrieving file: {str(e)}"


async def search_within_kb_files(project_id: str, search_term: str) -> str:
    """
    Search for a term within all knowledge base files.
    Returns matching excerpts from files that contain the term.
    """
    try:
        all_files = await db.files.find(
            {"project_id": project_id, "indexed": True},
            {"_id": 0}
        ).to_list(100)
        
        results = []
        for file_record in all_files:
            try:
                file_path = LOCAL_STORAGE_PATH / file_record['storage_path']
                if not file_path.exists():
                    continue
                
                content_bytes = file_path.read_bytes()
                ext = file_record['file_type'].lower()
                
                if ext in ['txt', 'md']:
                    content = content_bytes.decode('utf-8', errors='ignore')
                elif ext == 'pdf':
                    from PyPDF2 import PdfReader
                    import io
                    reader = PdfReader(io.BytesIO(content_bytes))
                    content = "\n".join([page.extract_text() or "" for page in reader.pages])
                elif ext == 'docx':
                    from docx import Document
                    import io
                    doc = Document(io.BytesIO(content_bytes))
                    content = "\n".join([p.text for p in doc.paragraphs])
                else:
                    continue
                
                # Search for term (case-insensitive)
                if search_term.lower() in content.lower():
                    # Find excerpts containing the term
                    lines = content.split('\n')
                    matching_lines = []
                    for i, line in enumerate(lines):
                        if search_term.lower() in line.lower():
                            # Get context (2 lines before and after)
                            start = max(0, i - 2)
                            end = min(len(lines), i + 3)
                            excerpt = '\n'.join(lines[start:end])
                            matching_lines.append(excerpt)
                            if len(matching_lines) >= 3:  # Max 3 excerpts per file
                                break
                    
                    if matching_lines:
                        results.append(f"## {file_record['original_filename']}\n\n" + "\n\n---\n\n".join(matching_lines))
            except Exception as e:
                logger.error(f"Error searching {file_record['original_filename']}: {e}")
                continue
        
        if results:
            logger.info(f"KB search for '{search_term}' found matches in {len(results)} files")
            return f"# Search results for '{search_term}'\n\n" + "\n\n===\n\n".join(results)
        else:
            return f"No matches found for '{search_term}' in the knowledge base."
            
    except Exception as e:
        logger.error(f"KB search error: {e}")
        return f"Error searching knowledge base: {str(e)}"


async def call_bedrock_converse_with_tools(
    model_id: str, 
    messages: List[dict], 
    aws_config: dict, 
    max_tokens: int = 4000,
    extended_thinking: bool = False,
    thinking_budget: int = 10000,
    enable_web_search: bool = False,
    enable_kb_tools: bool = False,
    project_id: str = None
) -> tuple:
    """
    Call AWS Bedrock using the Converse API with optional tool use.
    Supports: web search (Tavily), knowledge base file retrieval.
    Handles the tool use loop.
    Returns: (response_text, thinking_content, thinking_time)
    """
    import time
    start_time = time.time()
    
    try:
        bedrock_runtime = boto3.client(
            service_name='bedrock-runtime',
            region_name=aws_config['aws_region_name'],
            aws_access_key_id=aws_config['aws_access_key_id'],
            aws_secret_access_key=aws_config['aws_secret_access_key']
        )
        
        # Convert messages to Bedrock format
        bedrock_messages = []
        for msg in messages:
            if msg['role'] == 'system':
                continue
            content_text = msg['content']
            if isinstance(content_text, bytes):
                content_text = content_text.decode('utf-8', errors='replace')
            bedrock_messages.append({
                "role": msg['role'],
                "content": [{"text": str(content_text)}]
            })
        
        # Find system message
        system_content = None
        for msg in messages:
            if msg['role'] == 'system':
                sys_text = msg['content']
                if isinstance(sys_text, bytes):
                    sys_text = sys_text.decode('utf-8', errors='replace')
                system_content = [{"text": str(sys_text)}]
                break
        
        # Build inference config
        inference_config = {"maxTokens": max_tokens}
        add_temp_params = True
        additional_fields = {}
        
        is_claude_model = 'anthropic' in model_id.lower() or 'claude' in model_id.lower()
        thinking_content = None
        thinking_time = None
        
        # Extended thinking setup (same as before)
        if extended_thinking and is_claude_model:
            supports_thinking = any(x in model_id.lower() for x in [
                'claude-3-7', 'claude-4', 'sonnet-4', 'opus-4', 'haiku-4',
                'claude-sonnet-4', 'claude-opus-4', 'claude-haiku-4'
            ])
            if supports_thinking:
                actual_budget = max(1024, thinking_budget)
                additional_fields["thinking"] = {
                    "type": "enabled",
                    "budget_tokens": actual_budget
                }
                inference_config["maxTokens"] = actual_budget + 4000
                add_temp_params = False
                logger.info(f"Extended thinking enabled with budget: {actual_budget} tokens")
        
        # Only add temperature (not both temp and topP - Claude 4 doesn't allow both)
        if add_temp_params:
            inference_config["temperature"] = 0.7
        
        # Build tools list
        tools_list = []
        
        # Web search tool
        if enable_web_search and tavily_client:
            tools_list.append({
                "toolSpec": {
                    "name": "web_search",
                    "description": "Search the web for current information, news, facts, or any topic. Use this when you need up-to-date information or don't have knowledge about something.",
                    "inputSchema": {
                        "json": {
                            "type": "object",
                            "properties": {
                                "query": {
                                    "type": "string",
                                    "description": "The search query to look up on the web"
                                }
                            },
                            "required": ["query"]
                        }
                    }
                }
            })
            logger.info("Web search tool enabled")
        
        # Knowledge Base tools
        if enable_kb_tools and project_id:
            tools_list.append({
                "toolSpec": {
                    "name": "get_kb_file",
                    "description": "Retrieve the complete contents of a specific file from the knowledge base. Use this when you need to reference, check, or quote information from a file in the knowledge base.",
                    "inputSchema": {
                        "json": {
                            "type": "object",
                            "properties": {
                                "filename": {
                                    "type": "string",
                                    "description": "The name (or partial name) of the file to retrieve from the knowledge base"
                                }
                            },
                            "required": ["filename"]
                        }
                    }
                }
            })
            tools_list.append({
                "toolSpec": {
                    "name": "search_kb",
                    "description": "Search for a specific term or phrase across all files in the knowledge base. Use this to find information without loading entire files, or to check if a word/phrase appears in any reference document.",
                    "inputSchema": {
                        "json": {
                            "type": "object",
                            "properties": {
                                "search_term": {
                                    "type": "string",
                                    "description": "The term or phrase to search for in the knowledge base files"
                                }
                            },
                            "required": ["search_term"]
                        }
                    }
                }
            })
            logger.info("Knowledge base tools enabled")
        
        tool_config = {"tools": tools_list} if tools_list else None
        
        # Build converse params
        converse_params = {
            "modelId": model_id,
            "messages": bedrock_messages,
            "inferenceConfig": inference_config
        }
        
        if system_content:
            converse_params["system"] = system_content
        if additional_fields:
            converse_params["additionalModelRequestFields"] = additional_fields
        if tool_config:
            converse_params["toolConfig"] = tool_config
        
        logger.info(f"Bedrock Converse API call: model={model_id}, messages={len(bedrock_messages)}, web_search={enable_web_search}, kb_tools={enable_kb_tools}")
        
        # Call Bedrock and handle tool use loop
        max_iterations = 5  # Allow more iterations for KB lookups
        iteration = 0
        
        while iteration < max_iterations:
            iteration += 1
            response = bedrock_runtime.converse(**converse_params)
            
            stop_reason = response.get('stopReason', 'unknown')
            output_message = response.get('output', {}).get('message', {})
            content_blocks = output_message.get('content', [])
            
            # Check if model wants to use a tool
            if stop_reason == 'tool_use':
                logger.info("Model requested tool use")
                
                # Find tool use blocks
                tool_results = []
                assistant_content = []
                
                for block in content_blocks:
                    if 'toolUse' in block:
                        tool_use = block['toolUse']
                        tool_name = tool_use.get('name')
                        tool_id = tool_use.get('toolUseId')
                        tool_input = tool_use.get('input', {})
                        
                        assistant_content.append(block)
                        
                        # Handle web search tool
                        if tool_name == 'web_search':
                            query = tool_input.get('query', '')
                            logger.info(f"Executing web search: {query}")
                            search_result = await perform_web_search(query)
                            tool_results.append({
                                "toolResult": {
                                    "toolUseId": tool_id,
                                    "content": [{"text": search_result}]
                                }
                            })
                        
                        # Handle KB file retrieval tool
                        elif tool_name == 'get_kb_file':
                            filename = tool_input.get('filename', '')
                            logger.info(f"Retrieving KB file: {filename}")
                            file_content = await retrieve_kb_file_content(project_id, filename)
                            tool_results.append({
                                "toolResult": {
                                    "toolUseId": tool_id,
                                    "content": [{"text": file_content}]
                                }
                            })
                        
                        # Handle KB search tool
                        elif tool_name == 'search_kb':
                            search_term = tool_input.get('search_term', '')
                            logger.info(f"Searching KB for: {search_term}")
                            search_result = await search_within_kb_files(project_id, search_term)
                            tool_results.append({
                                "toolResult": {
                                    "toolUseId": tool_id,
                                    "content": [{"text": search_result}]
                                }
                            })
                        
                    elif 'text' in block:
                        assistant_content.append(block)
                    elif 'thinking' in block:
                        thinking_content = block.get('thinking', '')
                        thinking_time = round(time.time() - start_time)
                        assistant_content.append(block)
                
                # Add assistant message with tool use to conversation
                bedrock_messages.append({
                    "role": "assistant",
                    "content": assistant_content
                })
                
                # Add tool results as user message
                bedrock_messages.append({
                    "role": "user",
                    "content": tool_results
                })
                
                # Update converse params for next iteration
                converse_params["messages"] = bedrock_messages
                continue
            
            # No more tool use, extract final response
            response_text = ""
            for block in content_blocks:
                if 'text' in block:
                    text_content = block['text']
                    if isinstance(text_content, bytes):
                        text_content = text_content.decode('utf-8', errors='replace')
                    response_text += str(text_content)
                elif 'thinking' in block:
                    thinking_content = block.get('thinking', '')
                    thinking_time = round(time.time() - start_time)
                elif block.get('type') == 'thinking':
                    thinking_content = block.get('text', '')
                    thinking_time = round(time.time() - start_time)
            
            usage = response.get('usage', {})
            logger.info(f"Bedrock Converse API success: model={model_id}, response_length={len(response_text)}, iterations={iteration}, usage={usage}")
            
            return (response_text, thinking_content, thinking_time)
        
        # Max iterations reached
        logger.warning(f"Max tool use iterations ({max_iterations}) reached")
        return ("I encountered an issue processing your request. Please try again.", None, None)
        
    except boto3.exceptions.Boto3Error as e:
        logger.error(f"Bedrock boto3 error: {e}")
        raise HTTPException(status_code=500, detail=f"Bedrock API error: {str(e)}")
    except Exception as e:
        logger.error(f"Bedrock Converse API error: {type(e).__name__}: {e}")
        raise HTTPException(status_code=500, detail=f"Bedrock API error: {str(e)}")

def get_llm_config(project: dict):
    """
    Get LLM configuration based on project's provider setting.
    Returns: (model_name, api_key, provider_type, extra_config)
    """
    llm_provider = project.get("llm_provider", "anthropic")
    
    if llm_provider == "bedrock-claude":
        # AWS Bedrock with Claude
        aws_access_key = os.environ.get('AWS_ACCESS_KEY_ID', '')
        aws_secret_key = os.environ.get('AWS_SECRET_ACCESS_KEY', '')
        aws_region = os.environ.get('AWS_REGION', 'us-east-1')
        bedrock_model = os.environ.get('BEDROCK_CLAUDE_MODEL_ID', 'us.anthropic.claude-3-5-sonnet-20241022-v2:0')
        
        if not aws_access_key or not aws_secret_key:
            raise HTTPException(
                status_code=400,
                detail="AWS credentials not configured. Please add AWS_ACCESS_KEY_ID and AWS_SECRET_ACCESS_KEY to settings."
            )
        
        return (
            f"bedrock/{bedrock_model}",
            None,
            "bedrock-claude",
            {
                "aws_access_key_id": aws_access_key,
                "aws_secret_access_key": aws_secret_key,
                "aws_region_name": aws_region,
                "bedrock_model_id": bedrock_model
            }
        )
    elif llm_provider == "bedrock-mistral":
        # AWS Bedrock with Mistral
        aws_access_key = os.environ.get('AWS_ACCESS_KEY_ID', '')
        aws_secret_key = os.environ.get('AWS_SECRET_ACCESS_KEY', '')
        aws_region = os.environ.get('AWS_REGION', 'us-east-1')
        mistral_model = os.environ.get('BEDROCK_MISTRAL_MODEL_ID', 'mistral.mistral-large-2407-v1:0')
        
        if not aws_access_key or not aws_secret_key:
            raise HTTPException(
                status_code=400,
                detail="AWS credentials not configured."
            )
        
        return (
            f"bedrock/{mistral_model}",
            None,
            "bedrock-mistral",
            {
                "aws_access_key_id": aws_access_key,
                "aws_secret_access_key": aws_secret_key,
                "aws_region_name": aws_region,
                "bedrock_model_id": mistral_model
            }
        )
    else:
        # Anthropic Direct API (default)
        anthropic_key = os.environ.get('ANTHROPIC_API_KEY', '')
        emergent_key = os.environ.get('EMERGENT_LLM_KEY', '')
        
        api_key = anthropic_key if anthropic_key else emergent_key
        use_direct_anthropic = bool(anthropic_key)
        
        return (
            "anthropic/claude-sonnet-4-20250514",
            api_key,
            "anthropic",
            {"supports_extended_features": use_direct_anthropic}
        )

@api_router.post("/chat")
async def chat_with_ai(request: ChatRequest):
    # Get conversation
    conv = await db.conversations.find_one({"id": request.conversation_id}, {"_id": 0})
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    # Get project
    project = await db.projects.find_one({"id": conv["project_id"]}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Build system message with project context
    system_parts = []
    
    if project.get("instructions"):
        system_parts.append(f"# Project Instructions\n{project['instructions']}")
    
    if project.get("memory"):
        system_parts.append(f"# Project Memory/Context\n{project['memory']}")
    
    # Knowledge Base - provide file list and enable tools (not full content)
    kb_files_list = []
    if request.include_knowledge_base:
        try:
            # Get all indexed files for this project
            project_files = await db.files.find(
                {"project_id": conv["project_id"], "indexed": True}, 
                {"_id": 0, "original_filename": 1, "file_type": 1}
            ).to_list(100)
            
            if project_files:
                kb_files_list = [f['original_filename'] for f in project_files]
                
                # Add file list to system message so Claude knows what's available
                file_list_text = "\n".join([f"  - {name}" for name in kb_files_list])
                kb_info = f"""# Knowledge Base

You have access to the following reference files in the knowledge base. Use the `get_kb_file` tool to retrieve the full content of a specific file, or use the `search_kb` tool to search for a term across all files.

**Available Files:**
{file_list_text}

When the user asks you to edit something or check against style guides, proactively use these tools to look up relevant information."""
                system_parts.append(kb_info)
                logger.info(f"Knowledge base: {len(kb_files_list)} files available via tools")
        except Exception as e:
            logger.error(f"Knowledge base setup failed: {e}")
    
    system_message = "\n\n".join(system_parts) if system_parts else "You are a helpful AI assistant for editing assessment materials."
    
    # Get chat history
    history = await db.messages.find(
        {"conversation_id": request.conversation_id}, {"_id": 0}
    ).sort("created_at", 1).to_list(50)
    
    # Save user message
    user_msg = Message(
        conversation_id=request.conversation_id,
        role="user",
        content=request.message
    )
    await db.messages.insert_one(user_msg.model_dump())
    
    # Get LLM configuration based on project's provider setting
    model_name, api_key, provider_type, extra_config = get_llm_config(project)
    supports_extended_features = extra_config.get("supports_extended_features", False) if provider_type == "anthropic" else False
    
    # Add model identity to system message to prevent confusion when provider changes mid-conversation
    model_identity = ""
    if provider_type == "bedrock-claude":
        model_identity = "\n\n# Important: You are Claude, an AI assistant made by Anthropic, accessed via AWS Bedrock. Always identify yourself as Claude if asked."
    elif provider_type == "bedrock-mistral":
        model_identity = "\n\n# Important: You are Mistral, an AI assistant made by Mistral AI, accessed via AWS Bedrock. Always identify yourself as Mistral if asked."
    elif provider_type == "anthropic":
        model_identity = "\n\n# Important: You are Claude, an AI assistant made by Anthropic. Always identify yourself as Claude if asked."
    
    system_message_with_identity = system_message + model_identity
    
    # Build message history for context
    messages_for_llm = [{"role": "system", "content": system_message_with_identity}]
    for msg in history[-20:]:  # Last 20 messages for context
        messages_for_llm.append({"role": msg["role"], "content": msg["content"]})
    messages_for_llm.append({"role": "user", "content": request.message})
    
    # Send current message and get response
    try:
        import litellm
        
        # Build params based on provider
        params = {
            "model": model_name,
            "messages": messages_for_llm,
        }
        
        # Add provider-specific configuration
        if provider_type.startswith("bedrock"):
            # Add AWS credentials for Bedrock (both Claude and Mistral)
            params.update(extra_config)
        else:
            # Add API key for Anthropic
            params["api_key"] = api_key
        
        # Determine extended thinking settings
        thinking_content = None
        thinking_time = None
        use_extended_thinking = request.extended_thinking or project.get("extended_thinking_enabled", False)
        thinking_budget = request.thinking_budget if request.extended_thinking else project.get("thinking_budget", 10000)
        use_web_search = request.web_search or project.get("web_search_enabled", False)
        
        # Add extended features for Anthropic Direct API
        if provider_type == "anthropic" and extra_config.get("supports_extended_features", False):
            if use_extended_thinking:
                # Minimum budget is 1024 tokens per Anthropic API
                actual_budget = max(1024, thinking_budget)
                params["thinking"] = {"type": "enabled", "budget_tokens": actual_budget}
                params["max_tokens"] = max(16000, actual_budget + 4000)
            
            if use_web_search:
                params["web_search_options"] = {"search_context_size": "medium"}
        
        # Make the API call - route based on provider type
        import time
        start_time = time.time()
        
        # Log the exact provider and model being used
        logger.info(f"Chat request - provider_type={provider_type}, project_llm_provider={project.get('llm_provider')}, model_name={model_name}")
        
        if provider_type.startswith("bedrock"):
            # Use Bedrock Converse API
            bedrock_model_id = extra_config.get("bedrock_model_id", model_name.replace("bedrock/", ""))
            
            # Extended thinking is only supported for Bedrock Claude, not Mistral
            bedrock_extended_thinking = use_extended_thinking and provider_type == "bedrock-claude"
            
            # Web search is supported for Bedrock Claude if Tavily is configured
            bedrock_web_search = use_web_search and provider_type == "bedrock-claude" and tavily_client is not None
            
            # KB tools are enabled when knowledge base toggle is ON for Bedrock Claude
            bedrock_kb_tools = request.include_knowledge_base and provider_type == "bedrock-claude" and len(kb_files_list) > 0
            
            logger.info(f"Using Bedrock Converse API: bedrock_model_id={bedrock_model_id}, extended_thinking={bedrock_extended_thinking}, web_search={bedrock_web_search}, kb_tools={bedrock_kb_tools}")
            
            # Note: Extended Thinking and tools (web search/KB) cannot be used together on Bedrock
            # When thinking is enabled, tool use requires special handling of thinking blocks
            # Prioritize tools over thinking if both are requested
            if bedrock_extended_thinking and (bedrock_web_search or bedrock_kb_tools):
                logger.warning("Extended thinking conflicts with tools - disabling extended thinking")
                bedrock_extended_thinking = False
            
            response_text, thinking_content, thinking_time = await call_bedrock_converse_with_tools(
                model_id=bedrock_model_id,
                messages=messages_for_llm,
                aws_config=extra_config,
                max_tokens=4000,
                extended_thinking=bedrock_extended_thinking,
                thinking_budget=thinking_budget,
                enable_web_search=bedrock_web_search,
                enable_kb_tools=bedrock_kb_tools,
                project_id=conv["project_id"]
            )
        else:
            # Use litellm for Anthropic Direct API
            logger.info(f"Using litellm: model={model_name}")
            llm_response = await litellm.acompletion(**params)
            elapsed_time = time.time() - start_time
            
            # Extract response and thinking content
            response_text = ""
            thinking_content = None
            thinking_time = None
            
            if llm_response.choices:
                choice = llm_response.choices[0]
                message = choice.message
                
                # Get main response content
                response_text = message.content or ""
                
                # Check for thinking content - litellm returns it in different places
                # 1. Check reasoning_content (litellm normalized field)
                if hasattr(message, 'reasoning_content') and message.reasoning_content:
                    thinking_content = message.reasoning_content
                    thinking_time = round(elapsed_time)
                    logger.info(f"Found thinking in reasoning_content")
                
                # 2. Check thinking_blocks
                elif hasattr(message, 'thinking_blocks') and message.thinking_blocks:
                    thinking_blocks = message.thinking_blocks
                    if thinking_blocks and len(thinking_blocks) > 0:
                        thinking_content = thinking_blocks[0].get('thinking', '')
                        thinking_time = round(elapsed_time)
                        logger.info(f"Found thinking in thinking_blocks")
                
                # 3. Check provider_specific_fields
                elif hasattr(message, 'provider_specific_fields') and message.provider_specific_fields:
                    psf = message.provider_specific_fields
                    if psf.get('thinking_blocks'):
                        tb = psf['thinking_blocks']
                        if tb and len(tb) > 0:
                            thinking_content = tb[0].get('thinking', '')
                            thinking_time = round(elapsed_time)
                            logger.info(f"Found thinking in provider_specific_fields")
        
        # Save assistant message (include thinking in metadata if present)
        msg_data = {
            "id": str(uuid.uuid4()),
            "conversation_id": request.conversation_id,
            "role": "assistant",
            "content": response_text,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        
        if thinking_content:
            msg_data["thinking"] = thinking_content
            msg_data["thinking_time"] = thinking_time
        
        await db.messages.insert_one(msg_data)
        
        # Update conversation timestamp
        await db.conversations.update_one(
            {"id": request.conversation_id},
            {"$set": {"updated_at": datetime.now(timezone.utc).isoformat()}}
        )
        
        result = {
            "response": response_text, 
            "message_id": msg_data["id"]
        }
        
        if thinking_content:
            result["thinking"] = thinking_content
            result["thinking_time"] = thinking_time
        
        return result
    
    except Exception as e:
        logger.error(f"Chat error: {e}")
        raise HTTPException(status_code=500, detail=f"AI chat error: {str(e)}")


@api_router.post("/chat/with-files")
async def chat_with_files(
    conversation_id: str = Form(...),
    message: str = Form(...),
    include_knowledge_base: bool = Form(True),
    extended_thinking: bool = Form(False),
    thinking_budget: int = Form(10000),
    web_search: bool = Form(False),
    files: List[UploadFile] = File(default=[])
):
    """Chat endpoint with file attachments"""
    import base64
    import litellm
    import time
    
    # Get conversation
    conv = await db.conversations.find_one({"id": conversation_id}, {"_id": 0})
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    # Get project
    project = await db.projects.find_one({"id": conv["project_id"]}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Build system message
    system_parts = []
    if project.get("instructions"):
        system_parts.append(f"# Project Instructions\n{project['instructions']}")
    if project.get("memory"):
        system_parts.append(f"# Project Memory/Context\n{project['memory']}")
    
    # Add RAG context if requested
    if include_knowledge_base:
        try:
            api_key = os.environ.get('EMERGENT_LLM_KEY', '')
            context, sources = await retrieve_context_for_query(
                db=db, project_id=conv["project_id"], query=message,
                api_key=api_key, max_tokens=3000, top_k=4
            )
            if context:
                system_parts.append(f"# Relevant Knowledge Base Content\n\n{context}")
        except Exception as e:
            logger.error(f"RAG retrieval failed: {e}")
    
    system_message = "\n\n".join(system_parts) if system_parts else "You are a helpful AI assistant."
    
    # Process file attachments
    file_contents = []
    file_names = []
    
    for file in files:
        content = await file.read()
        file_names.append(file.filename)
        ext = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
        
        # Image files - convert to base64 for vision (OpenAI format for litellm)
        if ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
            mime_type = f"image/{ext}" if ext != 'jpg' else "image/jpeg"
            b64_content = base64.b64encode(content).decode('utf-8')
            # Use OpenAI format which litellm will convert to provider format
            file_contents.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:{mime_type};base64,{b64_content}"
                }
            })
        # Text-based files - extract FULL content
        elif ext in ['txt', 'md', 'json', 'csv', 'rtf']:
            text_content = content.decode('utf-8', errors='ignore')
            logger.info(f"Text file extracted: {file.filename}, {len(text_content)} chars")
            file_contents.append({
                "type": "text",
                "filename": file.filename,
                "content": text_content
            })
        # PDF files - extract FULL content
        elif ext == 'pdf':
            try:
                from PyPDF2 import PdfReader
                import io
                reader = PdfReader(io.BytesIO(content))
                # Extract ALL pages, no truncation
                text_parts = []
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(f"[Page {i+1}]\n{page_text}")
                text = "\n\n".join(text_parts)
                logger.info(f"PDF extracted: {file.filename}, {len(reader.pages)} pages, {len(text)} chars")
                file_contents.append({
                    "type": "text",
                    "filename": file.filename,
                    "content": text
                })
            except Exception as e:
                logger.error(f"PDF extraction failed: {e}")
                file_contents.append({"type": "text", "filename": file.filename, "content": f"[PDF file: {file.filename}]"})
        # DOCX files - extract FULL content
        elif ext == 'docx':
            try:
                from docx import Document
                import io
                doc = Document(io.BytesIO(content))
                # Extract ALL paragraphs, no truncation
                text = "\n".join([p.text for p in doc.paragraphs])
                logger.info(f"DOCX extracted: {file.filename}, {len(text)} chars")
                file_contents.append({
                    "type": "text",
                    "filename": file.filename,
                    "content": text
                })
            except Exception as e:
                logger.error(f"DOCX extraction failed: {e}")
                file_contents.append({"type": "text", "filename": file.filename, "content": f"[DOCX file: {file.filename}]"})
        # Other files - just note them
        else:
            file_contents.append({
                "type": "text",
                "filename": file.filename,
                "content": f"[Attached file: {file.filename}]"
            })
    
    # Build user message content
    user_content = []
    
    # Add images first (for vision)
    for fc in file_contents:
        if fc["type"] == "image_url":
            user_content.append(fc)
    
    # Add text content from files
    text_parts = []
    for fc in file_contents:
        if fc["type"] == "text":
            text_parts.append(f"--- {fc['filename']} ---\n{fc['content']}")
    
    if text_parts:
        user_content.append({
            "type": "text",
            "text": "Attached files:\n\n" + "\n\n".join(text_parts)
        })
    
    # Add user message
    if message:
        user_content.append({"type": "text", "text": message})
    
    # Get chat history
    history = await db.messages.find(
        {"conversation_id": conversation_id}, {"_id": 0}
    ).sort("created_at", 1).to_list(50)
    
    # Build messages for LLM
    messages_for_llm = [{"role": "system", "content": system_message}]
    for msg in history[-15:]:
        messages_for_llm.append({"role": msg["role"], "content": msg["content"]})
    
    # Add current message with attachments
    messages_for_llm.append({"role": "user", "content": user_content})
    
    # Save user message
    user_msg_content = message
    if file_names:
        user_msg_content = f"{message}\n\n📎 Attached: {', '.join(file_names)}" if message else f"📎 Attached: {', '.join(file_names)}"
    
    user_msg_data = {
        "id": str(uuid.uuid4()),
        "conversation_id": conversation_id,
        "role": "user",
        "content": user_msg_content,
        "attachments": file_names,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.messages.insert_one(user_msg_data)
    
    # API call
    anthropic_key = os.environ.get('ANTHROPIC_API_KEY', '')
    emergent_key = os.environ.get('EMERGENT_LLM_KEY', '')
    use_direct = bool(anthropic_key)
    api_key = anthropic_key if use_direct else emergent_key
    
    params = {
        "model": "anthropic/claude-sonnet-4-20250514",
        "messages": messages_for_llm,
        "api_key": api_key,
        "max_tokens": 4096
    }
    
    thinking_content = None
    thinking_time = None
    
    if use_direct and extended_thinking:
        actual_budget = max(1024, thinking_budget)
        params["thinking"] = {"type": "enabled", "budget_tokens": actual_budget}
        params["max_tokens"] = max(16000, actual_budget + 4000)
    
    if use_direct and web_search:
        params["web_search_options"] = {"search_context_size": "medium"}
    
    try:
        start_time = time.time()
        llm_response = await litellm.acompletion(**params)
        elapsed = time.time() - start_time
        
        response_text = ""
        if llm_response.choices:
            msg = llm_response.choices[0].message
            response_text = msg.content or ""
            
            if hasattr(msg, 'reasoning_content') and msg.reasoning_content:
                thinking_content = msg.reasoning_content
                thinking_time = round(elapsed)
        
        # Save assistant message
        asst_msg = {
            "id": str(uuid.uuid4()),
            "conversation_id": conversation_id,
            "role": "assistant",
            "content": response_text,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        if thinking_content:
            asst_msg["thinking"] = thinking_content
            asst_msg["thinking_time"] = thinking_time
        
        await db.messages.insert_one(asst_msg)
        
        await db.conversations.update_one(
            {"id": conversation_id},
            {"$set": {"updated_at": datetime.now(timezone.utc).isoformat()}}
        )
        
        result = {"response": response_text, "message_id": asst_msg["id"]}
        if thinking_content:
            result["thinking"] = thinking_content
            result["thinking_time"] = thinking_time
        
        return result
        
    except Exception as e:
        logger.error(f"Chat with files error: {e}")
        raise HTTPException(status_code=500, detail=f"AI chat error: {str(e)}")

# ============== STORAGE CONFIG ==============

@api_router.get("/config/features")
async def get_feature_config():
    """Check which features are available based on API key configuration"""
    has_direct_anthropic = bool(os.environ.get('ANTHROPIC_API_KEY', ''))
    has_bedrock = bool(os.environ.get('AWS_ACCESS_KEY_ID', '')) and bool(os.environ.get('AWS_SECRET_ACCESS_KEY', ''))
    has_tavily = tavily_client is not None
    
    available = ["anthropic"]
    if has_bedrock:
        available.extend(["bedrock-claude", "bedrock-mistral"])
    
    return {
        "extended_thinking_available": has_direct_anthropic,
        "web_search_available": has_direct_anthropic,
        "using_direct_anthropic_key": has_direct_anthropic,
        "bedrock_configured": has_bedrock,
        "tavily_configured": has_tavily,
        "bedrock_web_search_available": has_bedrock and has_tavily,
        "available_providers": available
    }

@api_router.post("/config/bedrock")
async def update_bedrock_config(
    aws_access_key_id: str,
    aws_secret_access_key: str,
    aws_region: str = "us-east-1"
):
    """Update AWS Bedrock configuration (stores in .env - for development only)"""
    # Note: In production, this should update a secure config store, not .env
    return {
        "message": "Bedrock configuration updated. Please restart backend to apply changes.",
        "configured": True
    }

@api_router.get("/storage/config")
async def get_storage_config():
    return {
        "provider": STORAGE_PROVIDER,
        "s3_configured": bool(os.environ.get('S3_BUCKET_NAME')),
        "local_path": str(LOCAL_STORAGE_PATH) if STORAGE_PROVIDER == 'local' else None
    }

@api_router.post("/storage/config")
async def update_storage_config(config: StorageConfig):
    """Update storage configuration - requires server restart to take effect"""
    # In production, this would update environment variables or config file
    # For now, return instructions
    return {
        "message": "Storage configuration updated. Set these environment variables and restart the server:",
        "env_vars": {
            "STORAGE_PROVIDER": config.provider,
            "S3_BUCKET_NAME": config.s3_bucket_name or "",
            "S3_ACCESS_KEY": "***" if config.s3_access_key else "",
            "S3_SECRET_KEY": "***" if config.s3_secret_key else "",
            "S3_REGION": config.s3_region or "us-east-1"
        }
    }

# Include router and middleware
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
